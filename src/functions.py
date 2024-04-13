# LIBRARY / MODULE / PUSTAKA

import streamlit as st
import pandas as pd
from docx import Document
import cv2
import os, math, random, re
from decimal import Decimal

from warnings import simplefilter

simplefilter(action= "ignore", category= FutureWarning)

# DEFAULT FUNCTIONS

"""Make Space

Fungsi-fungsi untuk membuat jarak pada webpage menggunakan margin space dengan
ukuran yang bervariatif
"""

def ms_20():
    st.markdown("<div class= \"ms-20\"></div>", unsafe_allow_html= True)

def ms_40():
    st.markdown("<div class= \"ms-40\"></div>", unsafe_allow_html= True)

def ms_60():
    st.markdown("<div class= \"ms-60\"></div>", unsafe_allow_html= True)

def ms_80():
    st.markdown("<div class= \"ms-80\"></div>", unsafe_allow_html= True)

"""Make Layout

Fungsi-fungsi untuk layouting wabpage menggunakan fungsi columns() dari
Streamlit.

Returns
-------
self : object containers
    Mengembalikan layout container.
"""

def ml_center():
    left, center, right = st.columns([0.3, 2.5, 0.3])
    return center

def ml_split():
    left, right = st.columns([1, 1])
    return left, right

def ml_left():
    left, right = st.columns([1.75, 1])
    return left, right

def ml_right():
    left, right = st.columns([1, 1.75])
    return left, right

"""Cetak text

Fungsi-fungsi untuk menampilkan teks dengan berbagai gaya menggunakan
method dari Streamlit seperti title(), write(), dan caption().

Parameters
----------
text : str
    Teks yang ingin ditampilkan dalam halaman.

size : int
    Ukuran Heading untuk teks yang akan ditampilkan.

underline : bool
    Kondisi yang menyatakan penambahan garis setelah teks ditampilkan.
"""

def show_title(text, underline= False):
    st.title(text)
    if underline:
        st.markdown("---")

def show_text(text, size= 3, underline= False):
    heading = "#" if size == 1 else (
        "##" if size == 2 else (
            "###" if size == 3 else (
                "####" if size == 4 else "#####"
            )
        )
    )

    st.write(f"{heading} {text}")
    if underline:
        st.markdown("---")

def show_caption(text, size= 3, underline= False):
    heading = "#" if size == 1 else (
        "##" if size == 2 else (
            "###" if size == 3 else (
                "####" if size == 4 else "#####"
            )
        )
    )

    st.caption(f"{heading} {text}")
    if underline:
        st.markdown("---")

def show_paragraf(text):
    st.markdown(f"<div><p>{text}</p></div>", unsafe_allow_html= True)

"""Load file

Fungsi-fungsi untuk membaca file dalam lokal direktori.

Parameters
----------
filepath : str
    Jalur tempat data tersedia di lokal direktori.

Returns
-------
self : object DataFrame or str
    Obyek dengan informasi yang berhasil diekstrak.
"""

def get_csv(filepath):
    return pd.read_csv(filepath)

def get_excel(filepath):
    return pd.read_excel(filepath)

def get_docx(filepath):
    doc = Document(filepath)
    text = []

    for row in doc.paragraphs:
        text.append(row.text)
    return "\n".join(text)

def to_docx(text: str, filename, filepath= "./data/keys"):
    """Simpan ke dalam file docx

    Parameters
    ----------
    text : str
        Teks yang akan ditulis ke dalam dokumen docx.

    filename : str
        Nama file.

    filepath : str
        Jalur direktori tempat dokumen akan disimpan.
    """
    mk_dir(filepath)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(f"{filepath}/{filename}")

def mk_dir(dirpath):
    """Buat folder

    Fungsi ini akan memeriksa path folder. Jika tidak ada folder yang di maksud
    dalam path tersebut, maka folder akan dibuat.

    Parameters
    ----------
    dirpath : str
        Jalur tempat folder akan dibuat.
    """
    if not os.path.exists(dirpath):
        os.makedirs(dirpath)

def duration_count(start, finish, msg= "berlalu"):
    """Duration count

    Menghitung durasi untuk setiap proses yang dilakukan dari awal hingga akhir

    Parameters
    ----------
    start : float
        Waktu mulai (dalam detik).

    finish : float
        Waktu akhir / selesai (dalam detik).
    """
    duration = finish - start

    hours = duration // 3600
    minutes = (duration - (hours * 3600)) // 60
    seconds = duration - ((hours * 3600) + (minutes * 60))

    if hours >= 1:
        text = f"Waktu {msg} : {int(hours)} jam {int(minutes)} menit {seconds:.2f} detik"
    elif minutes >= 1:
        text = f"Waktu {msg} : {int(minutes)} menit {seconds:.2f} detik"
    else:
        text = f"Waktu {msg} : {seconds:.2f} detik"

    return text

# CUSTOM FUNCTIONS
        
def embed_msg(filepath: str, message: str):
    """Image based Steganography using Least Significant Bit

    Menyematkan pesan dalam gambar dengan teknik LSB (Least Significant Bit)
    dengan memanipulasi bit-bit piksel gambar tanpa mengubah tampilan visual
    secara signifikan.

    Parameters
    ----------
    filepath : str
        Jalur file gambar yang akan diolah.
    
    message : str
        Pesan yang akan disematkan ke dalam gambar.

    Returns
    -------
    image : NumPy array
        Mengembalikkan NumPy array yang mengandung nilai piksel dari gambar
        dalam format BGR (blue-green-red).
    """
    image = cv2.imread(filepath)

    message = [format(ord(i), "08b") for i in message]
    _, width, _ = image.shape

    pixel_req = len(message) * 3
    row_req = pixel_req / width
    row_req = math.ceil(row_req)

    count, char_count = 0, 0
    for i in range(row_req + 1):
        while count < width and char_count < len(message):
            char = message[char_count]
            char_count += 1
            for id, bit in enumerate(char):
                if (bit == "1" and image[i][count][id % 3] % 2 == 0) or \
                (bit == "0" and image[i][count][id % 3] % 2 == 1):
                    image[i][count][id % 3] -= 1

                if id % 3 == 2:
                    count += 1

                if id == 7:
                    if char_count * 3 < pixel_req and image[i][count][2] % 2 == 1:
                        image[i][count][2] -= 1
                    if char_count * 3 >= pixel_req and image[i][count][2] % 2 == 0:
                        image[i][count][2] -= 1
                    count += 1
        count = 0
    mk_dir("./assets/steno")
    cv2.imwrite("./assets/steno/result.png", image)
    image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    return image

def read_msg(filepath: str):
    """Extract hidden message

    Baca pesan yang disembunyikan dalam suatu gambar dengan ekstraksi teknik LSB.

    Parameters
    ----------
    filepath : str
        Jalur gambar yang memiliki pesan disematkan.

    Returns
    -------
    self : str
        Pesan yang berhasil diekstrak dari gambar.
    """
    image = cv2.imread(filepath)

    data, stop = [], False
    for row in image:
        row.tolist()
        for id, pixel in enumerate(row):
            if id % 3 == 2:
                data.append(bin(pixel[0])[-1])
                data.append(bin(pixel[1])[-1])

                if bin(pixel[2])[-1] == "1":
                    stop = True
                    break
            else:
                data.append(bin(pixel[0])[-1])
                data.append(bin(pixel[1])[-1])
                data.append(bin(pixel[2])[-1])
        if stop:
            break

    message = []
    for i in range(int((len(data) + 1) / 8)):
        message.append(data[i * 8:(i * 8 + 8)])
    message = [chr(int("".join(i), 2)) for i in message]
    return "".join(message)

def to_ascii(text):
    """Mengonversi pesan teks ke dalam bentuk ASCII.

    Parameters
    ----------
    text : str
        Teks atau pesan yang akan diolah.

    Returns
    -------
    self : list
        Daftar yang berisi nilai ASCII setiap karakter.
    """
    ascii_form = []
    for char in text:
        ascii_form.append(ord(char))
    return ascii_form

def convert_data_size(text, size_bit= None):
    """Konversi data menjadi ukuran tertentu dalam byte.

    Parameters
    ----------
    text : str
        Data atau pesan yang akan dikonversi.

    size_bit : int
        Ukuran target dalam bit.

    Return
    ------
    self : str
        Teks dalam format ASCII yang telah dikonversi ukuran bit-nya.
    """
    if size_bit is None:
        size_bit = len(text.encode("utf-8")) * 8
    
    thresh = int(size_bit / 8)

    while len(text.encode("utf-8")) < thresh:
        text += " "

    bit_text = text.encode("utf-8")[:thresh]
    return bit_text.decode("utf-8")

def prime_fermat(num, num_tests= 100):
    """Prime with Fermat Test

    Cek apakah suatu bilangan adalah bilangan prima secara probabilitas dengan
    Uji Fermat.

    Parameters
    ----------
    num : int
        Bilangan yang akan diperiksa

    num_tests : int
        Jumlah perulangan dalam uji coba Fermat yang dilakukan.

    Returns
    -------
    self : bool
        True jika bilangan adalah bilangan prima. False jika bukan.
    """
    if num <= 1:
        return False
    elif num <= 3:
        return True
    
    def fermat_test(rand_num):
        """Fermat Test

        Parameters
        ----------
        rand_num : int
            Bilangan bulat random
        
        Returns
        -------
        self : bool
            True jika hasil uji berhasil dan False jika gagal.
        """
        if pow(rand_num, num - 1, num) != 1:
            return False
        return True
    
    for i in range(num_tests):
        rand_num = random.randint(2, num - 2)
        if not fermat_test(rand_num):
            return False
    return True

def generate_prime(min_bit, max_bit):
    """Generate prime value

    Menghasilkan bilangan prima antara panjang bit minimum dan maksimum.

    Parameters
    ----------
    min_bit : int
        Panjang bit minimum.

    max_bit : int
        Panjang bit maksimum.

    Returns
    -------
    self : int
        Bilangan prima dengan antara panjang bit minimum dan maksimum.
    """
    while True:
        candidate = random.getrandbits(random.randint(min_bit, max_bit))
        if candidate % 2 == 0:
            candidate += 1
        if prime_fermat(candidate):
            return candidate

def generate_pqr(min_bit, max_bit):
    """Generate prime (p) and random number (q and r)

    Menghasilkan bilangan prima acak untuk p, q, dan r, dengan p < q * r.

    Parameters
    ----------
    min_bit : int
        Panjang bit minimum untuk p (bilangan prima).
    
    max_bit : int
        Panjang bit maksimum untuk p (bilangan prima).

    Returns
    -------
    self : tuple
        Menghasilkan nilai tuple untuk p (bilangan prima), q, dan r
        (bilangan acak).
    """
    p = generate_prime(min_bit, max_bit)
    q = random.getrandbits(random.randint(min_bit, max_bit))
    r = random.getrandbits(random.randint(min_bit, max_bit))

    while p >= q * r:
        p = generate_prime(min_bit, max_bit)
        q = random.getrandbits(random.randint(min_bit, max_bit))
        r = random.getrandbits(random.randint(min_bit, max_bit))
    return p, q, r

def expand_euclidean(a, b):
    """Euclidean expansion

    Perluasan dari algoritma Euclidean untuk menemukan pembagi persekutuan
    terbesar (PBB) dari dua bilangan bulat. Ini juga untuk menemukan dua
    bilangan bulat yang memenuhi persamaan `ax + by = PBB(a, b)`.
    
    Parameters
    ----------
    a : int
        Bilangan bulat pertama.

    b : int
        Bilangan bulat kedua.

    Returns:
    X1 : int
        Nilai koefisien X1 yang merupakan invers modular bilangan bulat
        pertama terhadap bilangan bulat kedua. Hal ini dibuktikan dengan
        persamaan `ax + by = PBB(a, b)`.
    """
    D1, D2 = a, b
    X1, X2 = 1, 0
    Y1, Y2 = 0, 1

    while D2 > 0:
        K = D1 // D2
        D3 = D1 - K * D2
        D1 = D2
        D2 = D3
        
        X3 = X1 - K * X2
        X1 = X2
        X2 = X3

        Y3 = Y1 - K * Y2
        Y1 = Y2
        Y2 = Y3
    return X1

def inverse_modular(num, mod):
    """Calculate of Inverse Modular

    Mencari invers modular dari bilangan bulat terhadap nilai modulus.
    Invers modular digunakan untuk melakukan perkalian modular.

    Parameters
    ----------
    num : int
        Bilangan yang ingin dicari invers modular.
    
    mod : int
        Bilangan modulusnya.

    Returns
    -------
    inv : int or None
        Invers modular bilangan bulat terhadap modulus (dalam batasan
        0 <= invers <= modulus). None jika invers tidak ditemukan
        (misalnya bilangan bulat dan modulus tidak coprime atau
        bilangan prima).
    """
    inv = None
    while inv is None:
        temp = expand_euclidean(num, mod)

        if temp < 0:
            temp += mod
        if num * temp % mod == 1 and 0 <= temp <= mod:
            inv = temp
    return inv

def generate_keys():
    """Generate keys
    
    Menghasilkan kunci kriptografi Vincent-Sathiyamoorthy RSA.

    Returns
    -------
    self : tuple
        Berisi bilangan prima (p), bilangan acak (q dan r), kombinasi bilangan
        prima yang dhitung dengan `(q * r) - p` (s), ekponen publik untuk
        enkripsi pesan (e), nilai kombinasi untuk perhitungan eksponen privat
        dihitung dengan `(pow(p, 2) * s) + q` (t), modulus kunci untuk enkripsi
        dan dekripsi pesan (n), inverse modular dari p modulo n (u), dan
        eksponen privat untuk dekripsi pesan yang dihitung menggunakan invers
        modular u dari p modulo n (d)
    """
    min_bit = 1024
    max_bit = 2048

    p, q, r = generate_pqr(min_bit, max_bit)
    s = (q * r) - p
    e = (p * s) + r
    t = (pow(p, 2) * s) + q
    n = ((e * t) - p) // s
    u = inverse_modular(p, n)
    d = u * t

    return p, q, r, s, e, t, n, u, d

def read_key(filepath):
    """Read keys

    Membaca kunci kriptografi Vincent-Sathiyamoorthy RSA dari file docx.

    Parameters
    ----------
    filepath : str
        Jalur direktori dari file yang akan di olah.
    
    Returns
    -------
    self : tuple
        Kunci-kunci yang dimiliki, terdiri dari p, q, r, s, e, t, n, u, d.

    Lebih lanjut `Generate Keys Vincent-Sathiyamoorthy RSA`.
    """
    text = get_docx(filepath)
    keys = [int(Decimal(x.group())) for x in re.finditer(r"\d+", text)]
    p, q, r, s, e, t, n, u, d = keys[0], keys[1], keys[2], keys[3], \
        keys[4], keys[5], keys[6], keys[7], keys[8]
    return p, q, r, s, e, t, n, u, d